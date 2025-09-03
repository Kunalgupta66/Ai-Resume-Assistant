from django.shortcuts import render, redirect
from django.http import HttpResponse 
from django.views.decorators.http import require_POST
from django.template.loader import render_to_string
import io, json
from django.contrib.auth.decorators import login_required
import google.generativeai as genai
from docx import Document
from weasyprint import HTML
from .forms import CustomUserCreationForm  
from collections import Counter
from .ai_service import generate_structured_resume                 
from django.contrib.auth import login

# ---------- Home ----------
def home(request):
    return render(request, "home.html")


def register(request):
    if request.method == "POST":
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect("home")  # go to homepage after signup
    else:
        form = CustomUserCreationForm()
    return render(request, "register.html", {"form": form})


# ---------- Step 1: AI Description ----------
@login_required
def ai_description(request):
    return render(request, "ai_description.html")


@require_POST
def ai_generate(request):
    description = request.POST.get("description", "").strip()
    data = generate_structured_resume(description)
    request.session["resume_data"] = data  # store for step 2
    return redirect("resume_form")


# ---------- Step 2: Editable Form ----------
@login_required
def resume_form(request):
    # load existing or empty shell
    data = request.session.get("resume_data") or {
        "name": "", "email": "", "phone": "", "location": "", "summary": "",
        "skills": [], "education": [], "experience": [], "projects": [],
        "certifications": [], "achievements": [], "languages": [], "interests": [],
        "links": {"github": "", "linkedin": "", "portfolio": ""}
    }

    if request.method == "POST":
        payload = request.POST.get("resume_json")
        try:
            data = json.loads(payload)
        except Exception:
            pass
        request.session["resume_data"] = data
        return redirect("resume_preview")

    return render(request, "resume_form_dynamic.html", {"data_json": json.dumps(data)})


# ---------- Step 3: Resume Preview ----------
def resume_preview(request):
    data = request.session.get("resume_data")
    if not data:
        return redirect("ai_description")
    return render(request, "resume_preview.html", {"d": data})


# ---------- Downloads ----------
def download_resume(request, format):
    # Get data from session (already saved during resume form step)
    data = request.session.get("resume_data")
    if not data:
        return HttpResponse("No resume data found.", status=400)

    filename = request.GET.get("filename", "resume")

    if format == "pdf":
        # Render your existing template
        html_string = render_to_string("resume_preview.html", {"d": data})
        response = HttpResponse(content_type="application/pdf")
        response['Content-Disposition'] = f'attachment; filename={filename}.pdf'
        # Convert to PDF (WeasyPrint)
        HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(response)
        return response

    elif format == "docx":
        doc = Document()
        doc.add_heading(data.get("name", "Resume"), 0)
        doc.add_paragraph(f"{data.get('email','')} | {data.get('phone','')} | {data.get('location','')}")
        if data.get("summary"):
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(data["summary"])
        if data.get("skills"):
            doc.add_heading("Skills", level=1)
            for s in data["skills"]:
                doc.add_paragraph(f"{s.get('title','')} — {s.get('level','')}", style="List Bullet")
        # (Add education, experience, etc. same way...)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response['Content-Disposition'] = f'attachment; filename={filename}.docx'
        return response

    elif format == "txt":
        content = f"Resume: {data.get('name','')}\n\nSummary:\n{data.get('summary','')}\n\nSkills:\n"
        for s in data.get("skills", []):
            content += f"- {s.get('title','')} ({s.get('level','')})\n"
        response = HttpResponse(content, content_type="text/plain")
        response['Content-Disposition'] = f'attachment; filename={filename}.txt'
        return response

    else:
        return HttpResponse("Invalid format", status=400)


def download_resume_docx(request):
    data = request.session.get("resume_data")
    if not data:
        return redirect("ai_description")

    doc = Document()
    doc.add_heading(data.get("name", ""), 0)
    doc.add_paragraph(f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('location', '')}")

    if data.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        doc.add_heading("Skills", level=1)
        for s in data["skills"]:
            doc.add_paragraph(f"{s.get('title', '')} — {s.get('level', '')}", style="List Bullet")

    if data.get("experience"):
        doc.add_heading("Experience", level=1)
        for e in data["experience"]:
            doc.add_paragraph(f"{e.get('title', '')} @ {e.get('company', '')} ({e.get('start', '')} - {e.get('end', '')})", style="List Bullet")
            doc.add_paragraph(e.get("description", ""))

    if data.get("projects"):
        doc.add_heading("Projects", level=1)
        for pr in data["projects"]:
            doc.add_paragraph(f"{pr.get('name', '')} [{pr.get('tech', '')}]", style="List Bullet")
            doc.add_paragraph(pr.get("description", ""))

    if data.get("education"):
        doc.add_heading("Education", level=1)
        for ed in data["education"]:
            doc.add_paragraph(f"{ed.get('degree', '')} — {ed.get('institute', '')} ({ed.get('start', '')} - {ed.get('end', '')}) | {ed.get('score', '')}", style="List Bullet")

    if data.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for c in data["certifications"]:
            doc.add_paragraph(f"{c.get('name', '')} — {c.get('issuer', '')} {c.get('year', '')}", style="List Bullet")

    if data.get("languages"):
        doc.add_heading("Languages", level=1)
        doc.add_paragraph(", ".join(data["languages"]))

    if data.get("interests"):
        doc.add_heading("Interests", level=1)
        doc.add_paragraph(", ".join(data["interests"]))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = 'attachment; filename="resume.docx"'
    return response

@login_required
def interview_qna(request):
    qa_text = None

    if request.method == "POST":
        job_role = request.POST.get("job_role")
        company = request.POST.get("company", "")

        data = request.session.get("resume_data", {})

        prompt = f"""
        Generate 5 interview questions and their best possible answers 
        for a candidate applying as {job_role}.
        The candidate’s resume is: {data}.
        Focus on technical + HR questions. 
        {"Tailor questions as if for " + company if company else ""}
        """

        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)

        # ✅ store raw text
        qa_text = response.text if response and response.text else "⚠️ No response received."

    return render(request, "interview_qna.html", {"qa_text": qa_text})

